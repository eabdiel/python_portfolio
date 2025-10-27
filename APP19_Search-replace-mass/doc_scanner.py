"""
/***********************************************************************
* Program: Document Scanner
* Version: 3.3
* Created by: Arthrex IT SAP COE GRM Team
* Developer: Edwin Rodriguez
* Date: 2025-10-02
*
* Description:
*   Scans a directory of documents (PDF, DOCX, DOC)
*   for a given search string, regex pattern, or Bullet Modes.
*   Results are saved into an Excel file.
*
* Features:
*   - String or regex text search inside documents
*   - Generic Bullet Mode:
*       * For any file (if search string blank → all files, else filename match)
*       * Extracts bullets, includes header if present
*   - PMT Mode (PDF + Word):
*       * Allows only PDF/DOCX/DOC in the selected folder; errors out otherwise
*       * PDF parsing:
*           - Header: lines starting at "Document Number" (next non-empty lines)
*           - Tasks: lines after "Please complete all work listed below for the P.M."
*           - Comments: lines after "Comments / Notes:" until "Following fields/sections..."
*           - Strips common bullet glyphs (e.g., "•", "–", "", "►", "-")
*       * DOCX parsing (kept): header text from header region (fallback first table),
*                              marker-based Tasks/Comments
*       * DOC (legacy): text via textract, parsed with same logic as PDF
*   - All results go into one sheet
*   - Progress bar with stop option
*   - Graceful error logging
***********************************************************************/
"""

import os, re, threading, tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openpyxl import Workbook
from docx import Document
import textract, PyPDF2


stop_flag = False

ALLOWED_DOC_EXTS = {".pdf", ".docx", ".doc"}

def sanitize_text(text: str) -> str:
    try:
        return text.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")
    except Exception:
        return "(unreadable)"

# ---------- Shared PMT text parser (for PDF + DOC via plain text) ----------

HEADER_START_RE = re.compile(r"\bDocument\s+Number\b", re.IGNORECASE)
PM_MARKER_RE    = re.compile(r"Please\s+complete\s+all\s+work\s+listed\s+below\s+for\s+the\s+P\.?M\.?", re.IGNORECASE)
COMMENTS_RE     = re.compile(r"Comments\s*/\s*Notes\s*:", re.IGNORECASE)
STOP_RE         = re.compile(r"(Following\s+(fields|sections)\s+are\s+not\s+required|Name\s*\(Print\)|Additional\s+Work\s+Required|EAM\s+Work\s+Order\s+Number|Work\s+Completed\s+By|Work\s+Accepted\s+By)", re.IGNORECASE)

BULLET_PREFIX_RE = re.compile(r'^[\s\W]*(?:•|\-|\–||►|»|\*|\u2022|\u2023|\u25CF|\u25CB|\u25AA|\u2043|\u2794)\s*')

def parse_pmt_from_text(raw_text: str):
    """
    Returns (header_text, tasks_list, comments_list)
    - header_text: single-line string
    - tasks_list/comments_list: list[str]
    """
    text = sanitize_text(raw_text or "")
    # Normalize Windows newlines and multiple blanks
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln is not None]  # keep empties for control, but filter when needed

    # ---- HEADER: start at "Document Number", collect subsequent non-empty lines
    header_text = "(no header text)"
    for i, line in enumerate(lines):
        if HEADER_START_RE.search(line or ""):
            header_parts = []
            # Collect up to a reasonable number of subsequent non-empty lines,
            # stopping if a known lower-page marker appears.
            j = i
            while j < len(lines):
                nxt = (lines[j] or "").strip()
                if nxt:
                    if (PM_MARKER_RE.search(nxt) or
                        STOP_RE.search(nxt) or
                        re.search(r"\bEquipment\s+ID\s+Number\b", nxt, re.IGNORECASE) or
                        re.search(r"\bRMF[-\s]", nxt, re.IGNORECASE) or
                        re.search(r"\bPage\s+\d+\s+of\s+\d+\b", nxt, re.IGNORECASE)):
                        break
                    header_parts.append(nxt)
                    # Heuristic: most PMT headers are 2–4 lines; keep it tight.
                    if len(header_parts) >= 4:
                        # If we already captured "Document Number..." + 2-3 lines, stop.
                        pass
                j += 1
                # If we hit an empty line, we'll continue to try to grab the next content line,
                # but break if we see obvious footer or section start in the next iterations.
                if j < len(lines) and not (lines[j] or "").strip():
                    # Look ahead one more line; if it's empty again, bail.
                    if j + 1 < len(lines) and not (lines[j + 1] or "").strip():
                        break
            if header_parts:
                header_text = " ".join(header_parts)
            break

    # ---- TASKS & COMMENTS
    tasks, comments = [], []
    state = None  # None | "tasks" | "comments"
    for ln in lines:
        up = (ln or "").upper().strip()
        if not up:
            continue

        if PM_MARKER_RE.search(ln):
            state = "tasks"
            continue
        if COMMENTS_RE.search(ln):
            state = "comments"
            continue
        if STOP_RE.search(ln):
            # Stop collecting on the first definitive stop marker.
            state = None
            break

        if state in ("tasks", "comments"):
            cleaned = BULLET_PREFIX_RE.sub("", ln).strip()
            if cleaned:
                if state == "tasks":
                    tasks.append(cleaned)
                else:
                    comments.append(cleaned)

    return header_text, tasks, comments

# ---------- PMT (DOCX) Extraction (kept; small robustness tweaks) ----------

def extract_pmt_docx(path):
    results = []
    try:
        doc = Document(path)

        # Header from the Word header region; fallback to first table if empty
        header_parts = []
        try:
            for p in doc.sections[0].header.paragraphs:
                t = sanitize_text(p.text).strip()
                if t:
                    header_parts.append(t)
        except Exception:
            pass

        if (not header_parts) and doc.tables:
            try:
                for row in doc.tables[0].rows:
                    for cell in row.cells:
                        txt = sanitize_text(cell.text).strip()
                        if txt:
                            header_parts.append(txt)
            except Exception:
                pass

        header_text = " ".join(header_parts) if header_parts else "(no header text)"

        # Tasks/Comments by markers (case-insensitive)
        tasks, comments = [], []
        state = None
        for para in doc.paragraphs:
            text = sanitize_text(para.text).strip()
            if not text:
                continue
            if PM_MARKER_RE.search(text):
                state = "tasks"; continue
            if COMMENTS_RE.search(text):
                state = "comments"; continue
            if STOP_RE.search(text):
                state = None; break

            if state in ("tasks", "comments"):
                cleaned = BULLET_PREFIX_RE.sub("", text).strip()
                if cleaned:
                    (tasks if state == "tasks" else comments).append(cleaned)

        file_name = os.path.basename(path)
        if tasks or comments:
            for i, line in enumerate(tasks, start=1):
                results.append((file_name, header_text, "Task", line, i))
            for i, line in enumerate(comments, start=1):
                results.append((file_name, header_text, "Comments", line, i))
        else:
            results.append((file_name, header_text, "(none)", "(no tasks/comments found)", ""))

    except Exception as e:
        results.append((os.path.basename(path), "(error)", "", "", str(e), ""))
    return results

# ---------- PMT (PDF) Extraction ----------

def extract_pmt_pdf(path):
    results = []
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join([(page.extract_text() or "") for page in reader.pages])
        header_text, tasks, comments = parse_pmt_from_text(text)
        file_name = os.path.basename(path)

        if tasks or comments:
            for i, line in enumerate(tasks, start=1):
                results.append((file_name, header_text, "Task", line, i))
            for i, line in enumerate(comments, start=1):
                results.append((file_name, header_text, "Comments", line, i))
        else:
            results.append((file_name, header_text, "(none)", "(no tasks/comments found)", ""))

    except Exception as e:
        results.append((os.path.basename(path), "(error)", "", "", str(e), ""))
    return results

# ---------- PMT (DOC legacy) Extraction via textract ----------

def extract_pmt_doc(path):
    results = []
    try:
        text = textract.process(path).decode("utf-8", errors="ignore")
        header_text, tasks, comments = parse_pmt_from_text(text)
        file_name = os.path.basename(path)

        if tasks or comments:
            for i, line in enumerate(tasks, start=1):
                results.append((file_name, header_text, "Task", line, i))
            for i, line in enumerate(comments, start=1):
                results.append((file_name, header_text, "Comments", line, i))
        else:
            results.append((file_name, header_text, "(none)", "(no tasks/comments found)", ""))

    except Exception as e:
        results.append((os.path.basename(path), "(error)", "", "", str(e), ""))
    return results

# ---------- Generic Bullet Mode (unchanged) ----------

def extract_generic_docx(path, search_string):
    results = []
    try:
        if search_string and search_string.lower() not in os.path.basename(path).lower():
            return results
        doc = Document(path)
        header = " ".join(
            sanitize_text(p.text) for p in doc.sections[0].header.paragraphs if p.text.strip()
        ) or "(none)"
        file_name = os.path.basename(path)
        found_bullet = False
        for i, para in enumerate(doc.paragraphs):
            text = sanitize_text(para.text).strip()
            if para.style.name.startswith("List") or re.match(r'^[\s\W]*(•|-|–|)', text):
                line = re.sub(r'^[\s\W]*(•|-|–|)', '', text).strip()
                if line:
                    results.append((file_name, header, line, i+1))
                    found_bullet = True
        if not found_bullet:
            results.append((file_name, header, "(no bullets found)", ""))
    except Exception as e:
        results.append((os.path.basename(path), "(error)", str(e), ""))
    return results

# ---------- Standard Search (non-bullet modes) ----------

def read_pdf(path, search_string, regex_mode, regex_pattern):
    results = []
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = sanitize_text(page.extract_text() or "")
                for line in text.splitlines():
                    if match_line(line, search_string, regex_mode, regex_pattern):
                        results.append((os.path.basename(path), f"Page {i+1}", line.strip()))
    except Exception:
        results.append((os.path.basename(path), "(error)", "File could not be read"))
    return results

def read_docx(path, search_string, regex_mode, regex_pattern):
    results = []
    try:
        doc = Document(path)
        header = " ".join(sanitize_text(p.text) for p in doc.sections[0].header.paragraphs if p.text.strip()) or "(none)"
        for para in doc.paragraphs:
            line = sanitize_text(para.text)
            if match_line(line, search_string, regex_mode, regex_pattern):
                results.append((os.path.basename(path), header, line.strip()))
    except Exception:
        results.append((os.path.basename(path), "(error)", "File could not be read"))
    return results

def read_doc(path, search_string, regex_mode, regex_pattern):
    results = []
    try:
        text = textract.process(path).decode("utf-8", errors="ignore")
        text = sanitize_text(text)
        for line in text.splitlines():
            if match_line(line, search_string, regex_mode, regex_pattern):
                results.append((os.path.basename(path), "(legacy doc)", line.strip()))
    except Exception:
        results.append((os.path.basename(path), "(error)", "File could not be read"))
    return results

def match_line(line, search_string, regex_mode, regex_pattern):
    if regex_mode:
        try:
            return re.search(regex_pattern, line, re.IGNORECASE)
        except re.error:
            return False
    else:
        return search_string.lower() in line.lower()

# ---------- Directory Scanner ----------

def scan_directory(folder, search_string, regex_mode, regex_pattern, include_subs, file_type, bullet_mode_generic, bullet_mode_pmt, progress):
    global stop_flag

    # Build full file list (respecting include_subs)
    all_files = []
    for root, dirs, fs in os.walk(folder):
        for f in fs:
            all_files.append(os.path.join(root, f))
        if not include_subs:
            break

    # PMT mode: enforce only PDF/DOCX/DOC in folder (per requirement)
    if bullet_mode_pmt:
        offenders = [p for p in all_files if os.path.splitext(p)[1].lower() not in ALLOWED_DOC_EXTS]
        if offenders:
            sample = "\n".join(os.path.basename(x) for x in offenders[:8])
            messagebox.showerror(
                "Invalid files for PMT mode",
                "PMT Mode requires the folder to contain only PDF/DOCX/DOC files.\n"
                "Please remove or move other files and try again.\n\nFirst few offending files:\n" + sample
            )
            return

    # Determine which files to process for the chosen file_type filter
    exts = [".pdf",".docx",".doc"] if file_type=="All" else [f".{file_type.lower()}"]
    files = [p for p in all_files if any(p.lower().endswith(ext) for ext in exts)]

    wb = Workbook()
    ws = wb.active

    if bullet_mode_pmt:
        ws.append(["File Name","Header","Section","Content","Line #"])
    elif bullet_mode_generic:
        ws.append(["File Name","Header","Bullet Text","Line #"])
    else:
        ws.append(["File","Header/Page","Line"])

    total, rows_written = max(len(files), 1), 0
    for idx, file in enumerate(files):
        if stop_flag:
            break

        res = []
        base = os.path.basename(file).upper()
        ext  = os.path.splitext(file)[1].lower()

        if bullet_mode_pmt and base.startswith("PMT"):
            if ext == ".pdf":
                res = extract_pmt_pdf(file)
            elif ext == ".docx":
                res = extract_pmt_docx(file)
            elif ext == ".doc":
                res = extract_pmt_doc(file)
            else:
                # Should not happen due to earlier enforcement; skip defensively
                pass

        elif bullet_mode_generic and ext == ".docx":
            res = extract_generic_docx(file, search_string)

        else:
            # Standard match mode
            if ext == ".pdf":
                res = read_pdf(file, search_string, regex_mode, regex_pattern)
            elif ext == ".docx":
                res = read_docx(file, search_string, regex_mode, regex_pattern)
            elif ext == ".doc":
                res = read_doc(file, search_string, regex_mode, regex_pattern)

        for r in res:
            ws.append(r)
            rows_written += 1

        progress["value"] = (idx + 1) / total * 100
        progress.update_idletasks()

    wb.save("Scan_Results.xlsx")
    if rows_written == 0:
        messagebox.showwarning("No Matches","No results found.")
    else:
        messagebox.showinfo("Done",f"Scan completed! {rows_written} rows written to Scan_Results.xlsx")

# ---------- Threading & GUI ----------

def start_scan(folder, search_string, regex_mode, regex_pattern, include_subs, file_type, bullet_mode_generic, bullet_mode_pmt, progress):
    global stop_flag
    stop_flag = False
    thread = threading.Thread(
        target=scan_directory,
        args=(folder, search_string, regex_mode, regex_pattern, include_subs, file_type, bullet_mode_generic, bullet_mode_pmt, progress)
    )
    thread.start()

def stop_scan():
    global stop_flag
    stop_flag = True

def toggle_regex():
    if regex_var.get():
        regex_entry.config(state="normal")
        if not regex_entry.get().strip():
            s = search_entry.get().strip()
            if s:
                regex_entry.delete(0, tk.END)
                regex_entry.insert(0, re.escape(s))
    else:
        regex_entry.config(state="disabled")
        regex_entry.delete(0, tk.END)

def main():
    global search_entry, regex_var, regex_entry
    root = tk.Tk()
    root.title("Document Scanner")

    tk.Label(root, text="Search String:").pack()
    search_entry = tk.Entry(root, width=40); search_entry.pack()

    regex_var = tk.BooleanVar()
    tk.Checkbutton(root, text="Search by Regex", variable=regex_var, command=toggle_regex).pack()
    tk.Label(root, text="Regex Pattern (if enabled):").pack()
    regex_entry = tk.Entry(root, width=40, state="disabled"); regex_entry.pack()
    tk.Label(root, text="(Uses Python regex. Default auto-escapes your string; edit freely)", fg="gray", font=("Arial",8)).pack()

    bullet_var_generic = tk.BooleanVar()
    bullet_var_pmt = tk.BooleanVar()
    tk.Checkbutton(root, text="List Bullets (Generic)", variable=bullet_var_generic).pack()
    tk.Checkbutton(root, text="PMT Mode", variable=bullet_var_pmt).pack()

    tk.Label(root, text="Select Folder:").pack()
    folder_var = tk.StringVar()
    tk.Entry(root, textvariable=folder_var, width=40).pack()
    tk.Button(root, text="Browse", command=lambda: folder_var.set(filedialog.askdirectory())).pack()

    include_subs = tk.BooleanVar()
    tk.Checkbutton(root, text="Include Subfolders", variable=include_subs).pack()

    file_type = tk.StringVar(value="All")
    tk.Label(root, text="File Types:").pack()
    ttk.Combobox(root, textvariable=file_type, values=["All","PDF","DOCX","DOC"]).pack()

    progress = ttk.Progressbar(root, orient="horizontal", length=300, mode="determinate"); progress.pack()

    tk.Button(
        root, text="Start",
        command=lambda: start_scan(
            folder_var.get(), search_entry.get(), regex_var.get(), regex_entry.get(),
            include_subs.get(), file_type.get(), bullet_var_generic.get(), bullet_var_pmt.get(), progress
        )
    ).pack()
    tk.Button(root, text="Stop", command=stop_scan).pack()

    root.mainloop()

if __name__ == "__main__":
    main()
